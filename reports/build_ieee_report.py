from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "reports" / "BTSecBench_v2_IEEE_Report.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, size=8.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)


def style_table(table, header=True):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for r, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
            if header and r == 0:
                set_cell_shading(cell, "E8EEF5")


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    style_table(table)
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)
        if widths:
            table.rows[0].cells[i].width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if widths:
                cells[i].width = Inches(widths[i])
    return table


def add_paragraph(doc, text, style=None, align=None, bold_first=False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.03
    if align is not None:
        p.alignment = align
    if bold_first and ":" in text:
        first, rest = text.split(":", 1)
        r = p.add_run(first + ":")
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(9.5)
        r._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        r._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        p.add_run(rest)
    else:
        p.add_run(text)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(9.5)
        run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(12 if level == 1 else 10.5)
    run.font.color.rgb = RGBColor(31, 77, 120)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        r.font.name = "Times New Roman"
        r.font.size = Pt(9.2)
        r._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        r._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")


def add_figure(doc, rel_path, caption, width=5.8):
    path = ROOT / rel_path
    if not path.exists():
        add_paragraph(doc, f"Figure omitted because {rel_path} was not present.")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(2)
    run = cap.add_run(caption)
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(8.5)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")


def add_page(doc, title, paragraphs=None, bullets=None, table=None, figure=None):
    doc.add_page_break()
    add_heading(doc, title)
    for text in paragraphs or []:
        add_paragraph(doc, text)
    if bullets:
        add_bullets(doc, bullets)
    if table:
        add_table(doc, table["headers"], table["rows"], table.get("widths"))
    if figure:
        add_figure(doc, figure["path"], figure["caption"], figure.get("width", 5.8))


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(9.5)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("BTSecBench_v2 IEEE-Style Project Report")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    r = title.add_run("BTSecBench_v2: An IEEE-Style Analysis of a Backdoor Security Benchmark for Traffic Sign Recognition")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(15)
    r._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")

    add_paragraph(doc, "Prepared from the local project repository at C:\\Users\\Bittu\\OneDrive\\Desktop\\BTSecBench_v2.", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "IEEE-style project analysis report", align=WD_ALIGN_PARAGRAPH.CENTER)

    pages = [
        ("I. Introduction and Project Scope",
         ["Modern image classifiers are increasingly deployed in safety-relevant settings, where model behavior under malicious data manipulation is as important as ordinary clean-set accuracy. BTSecBench_v2 targets this problem by collecting training, attack, defense, evaluation, visualization, and reporting utilities in one repository. Its primary domain is traffic sign recognition using GTSRB-style images resized to 64 by 64 pixels across 43 classes.",
          "The project is best understood as an experimental workbench. It contains multiple attack classes, model backbones, data validators, training loops, explainability modules, benchmark runners, and dashboards. The strongest implemented thread is the backdoor workflow: generate or load clean splits, train or load a model, inject attack triggers, estimate attack success, run a defense, and export structured results for review."],
         ["The repository contains empty or minimal top-level documentation files, so this report derives project behavior from source code, configuration files, and generated outputs.",
          "The analysis emphasizes implemented behavior over intended behavior, which is important for engineering review and reproducibility."], None, None),
        ("II. Repository Inventory",
         ["The workspace is organized around conventional machine-learning concerns. The `attacks` package defines the poisoning interface and concrete triggers. The `models` package exposes CNN, ResNet18, MobileNetV3, and EfficientNet-B0 backbones through a registry. The `data` package contains dataset, split, transform, validator, cache, and statistics utilities. The `engine` package provides reusable training, checkpointing, scheduler, history, loss, evaluator, and early-stopping components.",
          "Higher-level experiment code lives in `analysis`, `benchmark`, `defenses`, `explainability`, `visualization`, and `dashboard`. The reports folder already contains JSON, CSV, Markdown, spreadsheet, and figure outputs, which makes the project unusually well-positioned for generating research reports even though the README and pyproject files do not provide narrative guidance."],
         None,
         {"headers": ["Area", "Representative files", "Role"], "rows": [
             ["Attacks", "badnets.py, blend.py, sig.py, wanet.py", "Trigger injection and label remapping"],
             ["Defenses", "strip.py, fine_tuning.py", "Detection and mitigation workflows"],
             ["Models", "cnn.py, resnet18.py, mobilenet_v3.py, efficientnet_b0.py", "Classifier backbones"],
             ["Reports", "benchmark.json, strip_final_results.json, figures", "Evidence exports for analysis"],
         ], "widths": [1.2, 2.7, 2.6]}, None),
        ("III. Problem Definition",
         ["The core security question is whether a traffic sign classifier can be trained, attacked, detected, and repaired under repeatable experimental settings. Backdoor attacks differ from evasion attacks because the malicious behavior is implanted during data or training-time compromise and triggered later by a specific visual pattern or transformation. The clean distribution should remain accurate while triggered inputs are redirected to a target class.",
          "BTSecBench_v2 frames this problem through target-class poisoning. The base attack class stores `target_class`, `poison_rate`, and `seed`, then exposes single-sample and batch poisoning methods. Concrete attacks preserve the interface while changing the trigger mechanism. This is a sensible design because downstream poisoned datasets and benchmarks can treat attacks as interchangeable components."],
         ["Security objective: measure clean accuracy, attack success rate, detection behavior, and post-defense recovery.",
          "Engineering objective: make attacks and model backbones swappable without rewriting training or evaluation logic.",
          "Research objective: produce comparable artifacts for model, attack, and defense combinations."], None, None),
        ("IV. Dataset and Preprocessing",
         ["The configured dataset is GTSRB with 43 traffic sign classes. The statistics file reports 39,209 images, average dimensions of approximately 50.84 by 50.33 pixels, and image-size variability from 25 pixels to more than 200 pixels in width or height. The project normalizes this variability by resizing images to 64 by 64 pixels through the transform pipeline.",
          "The dataset class reads CSV rows containing image paths and labels, opens each image with PIL, converts it to RGB, converts it to a NumPy array for Albumentations, and returns a dictionary with image, label, and path. This structure supports ordinary training, validation, explainability, backdoor injection, and STRIP reference-bank construction."],
         None, None,
         {"path": "data/statistics/class_distribution.png", "caption": "Fig. 1. Class-distribution artifact generated by the project.", "width": 5.2}),
        ("V. Data Quality and Split Handling",
         ["The configuration sets an 80/20 train-validation split, batch size 64, 40 nominal epochs, and deterministic seed 42. The data validators assert that expected GTSRB roots exist, that the training classes count is 43, that split CSVs exist, and that sampled image paths resolve. These checks are useful because missing files or silent path errors can invalidate a backdoor experiment before the model is ever trained.",
          "One implementation caveat is that several loader functions force `num_workers=0`, while the main configuration contains worker, prefetch, and persistent-worker options. This improves portability on Windows and notebook environments, but it means configuration values are not always reflected in runtime behavior. A future configuration audit should reconcile these settings."],
         ["Strength: CSV-based splits are inspectable and can be versioned.",
          "Strength: image paths are returned with samples, which helps debugging and explainability.",
          "Risk: README-level dataset acquisition instructions are absent.",
          "Risk: split provenance and seed lineage should be captured in experiment metadata."], None, None),
        ("VI. Model Architecture Layer",
         ["The model layer uses a registry pattern. `get_model` lowercases the requested model name, validates it against `MODEL_REGISTRY`, and constructs either the local CNN baseline or a transfer-learning backbone. Supported names include `cnn`, `resnet18`, `mobilenet_v3`, and `efficientnet_b0`. This is a pragmatic interface for benchmark automation because model selection becomes a configuration concern.",
          "The CNN baseline contains four convolutional blocks with batch normalization, ReLU, and max pooling, followed by a flattened classifier with a 512-unit hidden layer and dropout. EfficientNet-B0 replaces the final classifier layer with a 43-class linear head and optionally freezes the backbone. The benchmark output confirms all four models produce `(4, 43)` logits on `(4, 3, 64, 64)` inputs."],
         None,
         {"headers": ["Model", "Parameters", "Size", "CPU images/s"], "rows": [
             ["CNN", "2,509,099", "9.58 MB", "398.29"],
             ["ResNet18", "11,198,571", "42.76 MB", "168.88"],
             ["MobileNetV3", "4,257,115", "16.33 MB", "251.37"],
             ["EfficientNet-B0", "4,062,631", "15.66 MB", "176.50"],
         ], "widths": [1.6, 1.6, 1.3, 1.3]}, None),
        ("VII. Training Engine",
         ["The training engine is built around a reusable `Trainer` class. It moves batches to the selected device, computes loss and logits, clips gradients to a maximum norm of 5.0, steps the optimizer, updates a metric accumulator, prints progress through tqdm, validates through a separate evaluator, saves best and last checkpoints, updates training history, and supports early stopping.",
          "The configured optimizer is AdamW with learning rate 0.001 and weight decay 0.0001 for general training. The scheduler is CosineAnnealingLR. For fine-tuning defense, the learning rate is reduced to 1e-5 in the exported run, which is appropriate for mitigation where the model should recover clean behavior without destabilizing already learned representations."],
         ["Training loop reports accuracy, precision, recall, F1, balanced accuracy, MCC, learning rate, samples, epoch time, and throughput.",
          "History export to JSON and CSV supports downstream plotting and report generation.",
          "Checkpointing separates best model and last checkpoint, which is important for reproducibility and recovery."], None, None),
        ("VIII. Attack Framework",
         ["All backdoor attacks inherit from `BaseAttack`, which centralizes seed setting, poisoning probability, target-class relabeling, batch poisoning, and callable behavior. The concrete attack only needs to implement `apply_trigger`. This is the correct abstraction boundary because trigger logic changes across attacks while label policy and poisoning probability remain shared.",
          "BadNets applies a square trigger at configurable positions. Blend mixes every pixel with a trigger color using an alpha coefficient. SIG adds a sinusoidal signal across image columns. WaNet generates smooth displacement fields and warps images with OpenCV remapping. The implementations support NumPy images and torch tensors by converting tensors through HWC NumPy form and back to CHW tensor form."],
         None,
         {"headers": ["Attack", "Trigger mechanism", "Security interpretation"], "rows": [
             ["BadNets", "Visible square patch", "Classic localized trigger"],
             ["Blend", "Global alpha blend", "Distributed low-frequency trigger"],
             ["SIG", "Sinusoidal perturbation", "Pattern-based stealth trigger"],
             ["WaNet", "Smooth geometric warp", "Imperceptible spatial trigger"],
         ], "widths": [1.1, 2.4, 3.0]}, None),
        ("IX. BadNets Experiment",
         ["The available attack-success export reports model `efficientnet_b0`, attack `badnets`, target class 0, poison rate 1.0, and 7,842 validation samples. The recorded attack success rate is 33.28 percent. Because the export also reports accuracy as 33.28 percent with low precision and recall under the attack-success evaluation framing, the result should be read as partial target-class redirection under the chosen checkpoint and trigger settings rather than a universal result for BadNets.",
          "A 33.28 percent ASR is meaningful but not catastrophic by backdoor literature standards, where strong training-time poisoning often produces much higher target-class activation. The most likely explanations are checkpoint state, trigger visibility, target-class distribution, evaluation transform mismatch, poisoning setup, or insufficient backdoor-specific training. The report therefore treats this result as a baseline measurement, not the final attack capability of the framework."],
         None, None,
         {"path": "reports/figures/asr_comparison.png", "caption": "Fig. 2. Attack success rate comparison artifact.", "width": 5.2}),
        ("X. STRIP Defense Design",
         ["The STRIP detector is designed to identify poisoned inputs before inference. It builds a clean validation dataset, constructs a poisoned validation set using a selected attack, samples a clean reference bank, blends candidate images with random clean-bank samples, normalizes the perturbations, runs the model, and computes Shannon entropy over predicted probabilities. Low entropy under strong perturbation is treated as suspicious because a backdoor trigger can stabilize the target prediction.",
          "The implementation exposes model name, checkpoint, attack name, target class, poison rate, batch size, perturbation count, and blending alpha. Threshold optimization scans score values and selects the threshold that maximizes F1. This makes the exported results dependent on validation composition, poisoning strength, perturbation settings, and whether threshold selection is evaluated on the same distribution as final reporting."],
         ["Interpretability: entropy scores provide a simple scalar explanation for detector decisions.",
          "Risk: threshold optimization on the evaluation pool can overstate expected deployment behavior.",
          "Risk: a high false positive rate can make the defense unusable even when recall is high."], None, None),
        ("XI. STRIP Results",
         ["The final STRIP report records 20 perturbations, alpha 0.5, threshold 0.8742, accuracy 56.61 percent, precision 53.88 percent, recall 91.70 percent, F1 67.88 percent, false positive rate 78.49 percent, false negative rate 8.30 percent, and AUC 0.5278. The confusion matrix is [[1687, 6155], [651, 7191]], meaning the detector catches most poisoned samples but incorrectly flags many clean samples.",
          "This behavior is plausible for an aggressively tuned detector. High recall protects against missed triggered inputs, but false positives can suppress benign inferences, overload manual review, or reduce system availability. The AUC near 0.5 also indicates weak ranking separation between clean and poisoned entropy distributions under the tested settings."],
         None, None,
         {"path": "reports/figures/strip_metrics.png", "caption": "Fig. 3. STRIP metric summary generated by the benchmark.", "width": 5.2}),
        ("XII. STRIP Confusion Analysis",
         ["The confusion matrix reveals the operational character of the detector more clearly than accuracy alone. Of 7,842 clean samples, 6,155 are flagged as poisoned. Of 7,842 poisoned samples, 7,191 are correctly flagged and 651 are missed. This distribution produces the high recall and high false positive rate seen in the summary.",
          "For safety-critical traffic sign recognition, false negatives and false positives have different costs. A missed poisoned sign can cause targeted misclassification; a false positive can deny classification or trigger fallback behavior. The right threshold therefore depends on downstream control policy. If the benchmark is intended to compare defenses, it should report threshold-free curves and cost-weighted operating points in addition to the selected F1 threshold."],
         None, None,
         {"path": "reports/figures/strip_confusion_matrix.png", "caption": "Fig. 4. STRIP confusion matrix artifact.", "width": 5.0}),
        ("XIII. Fine-Tuning Defense",
         ["The fine-tuning defense loads a backdoored checkpoint, constructs clean train and validation loaders, optimizes with AdamW, applies a cosine scheduler, trains through the shared `Trainer`, saves a defended checkpoint, and exports JSON/CSV results. The exported configuration uses EfficientNet-B0, five epochs, learning rate 1e-5, and 7,842 validation samples.",
          "The resulting validation accuracy is 99.94 percent, validation loss 0.0040, precision 99.87 percent, recall 99.93 percent, F1 99.90 percent, balanced accuracy 99.93 percent, and MCC 0.9993. This shows excellent clean-set recovery. However, the exported fine-tuning result should be paired with post-defense attack-success evaluation before claiming true backdoor removal."],
         None,
         {"headers": ["Metric", "Value"], "rows": [
             ["Validation accuracy", "99.94%"],
             ["Validation loss", "0.0040"],
             ["Precision", "99.87%"],
             ["Recall", "99.93%"],
             ["F1 score", "99.90%"],
             ["MCC", "0.9993"],
         ], "widths": [2.4, 2.0]}, None),
        ("XIV. Training History Evidence",
         ["The exported five-epoch training history for the fine-tuning run shows stable high validation accuracy. Training loss declines from 0.1259 to 0.1040, while validation loss declines from 0.0116 to 0.0040. Validation accuracy remains around 99.95 percent across the run, peaking at 99.96 percent in epoch 4 and ending at 99.94 percent in epoch 5.",
          "This pattern suggests the model was already very strong on clean validation data before or early in the fine-tuning process. It also suggests that future defense reporting should not rely on clean validation accuracy alone. A robust mitigation report should include clean accuracy, attack success rate before defense, attack success rate after defense, and accuracy under non-targeted corruptions or adaptive triggers."],
         None, None,
         {"path": "reports/figures/accuracy_comparison.png", "caption": "Fig. 5. Accuracy comparison artifact from generated reports.", "width": 5.2}),
        ("XV. Benchmark Summary",
         ["The repository already aggregates selected results into a benchmark table. The table includes BadNets ASR, STRIP detection statistics, and fine-tuning precision/recall. This is useful because it demonstrates an end-to-end evidence pipeline from experiment scripts into machine-readable artifacts and visual summaries.",
          "The current aggregation is still narrow. A mature benchmark should expand the matrix across all implemented attacks, all supported backbones, several poison rates, multiple seeds, clean accuracy before attack, backdoor accuracy, defense-after-attack measurements, and runtime costs. The benchmark runner files are a natural place to formalize this grid."],
         None, None,
         {"path": "reports/figures/benchmark_summary.png", "caption": "Fig. 6. Benchmark summary figure exported by the project.", "width": 5.2}),
        ("XVI. Explainability Components",
         ["The repository contains Grad-CAM, Grad-CAM++, saliency, occlusion, and integrated-gradient modules. Their presence is important because backdoor research benefits from explanations that show whether a model attends to trigger regions, class-specific traffic-sign structure, or irrelevant background. Explainability is also useful for debugging failed attacks and defenses.",
          "The reportable opportunity is to align explainability outputs with the attack matrix. For each model-attack pair, the benchmark could generate clean and poisoned saliency maps, compare attribution concentration near trigger areas, and summarize whether fine-tuning reduces trigger-region attribution. This would turn explainability from an auxiliary feature into a validation tool for security claims."],
         ["Use Grad-CAM-like methods for spatial trigger visibility.",
          "Use integrated gradients for input-level contribution analysis.",
          "Use occlusion tests to quantify sensitivity to patches and warped regions.",
          "Export representative figures alongside numeric attack and defense metrics."], None, None),
        ("XVII. Dashboard and Visualization Layer",
         ["The visualization package contains radar, plot, heatmap, confusion, and dashboard modules, and the reports directory includes multiple PNG figures. Even though some dashboard files appear minimal in this snapshot, the project already treats reporting as a first-class workflow. This matters because security benchmarks are only useful when their outputs are easy to compare and audit.",
          "A future dashboard should expose the benchmark matrix as filters: model, attack, poison rate, target class, defense, seed, and checkpoint. It should also show confusion matrices, ROC curves, entropy distributions, attack-success bars, clean accuracy, runtime, and links to exact JSON artifacts. Such a dashboard would reduce the chance that users interpret a single metric out of context."],
         None, None,
         {"path": "reports/figures/strip_roc_curve.png", "caption": "Fig. 7. STRIP ROC curve artifact.", "width": 5.2}),
        ("XVIII. Software Engineering Assessment",
         ["The project shows good separation of concerns. Attacks, models, data, training, evaluation, analysis, reporting, and tracking are separated into packages. Registry and factory patterns reduce conditional logic in experiment scripts. JSON/CSV exports support reproducibility and downstream analysis. The engine layer records a rich metric set rather than accuracy alone.",
          "The main engineering gaps are documentation, test completeness, configuration consistency, and style hygiene. Several files are empty or placeholder-like. Some source comments contain encoding artifacts. The README and pyproject do not document installation, dataset preparation, or command usage. Test files are present but mostly empty or shallow in the current workspace. These gaps are repairable and do not undermine the core prototype, but they limit external reproducibility."],
         None,
         {"headers": ["Strength", "Evidence"], "rows": [
             ["Modularity", "Separate packages for attacks, data, models, engine, defenses, analysis"],
             ["Reproducibility artifacts", "JSON, CSV, Markdown, XLSX, and PNG reports exist"],
             ["Extensibility", "Factories and registries for attacks and models"],
             ["Metric breadth", "Accuracy, precision, recall, F1, balanced accuracy, MCC, AUC"],
         ], "widths": [1.8, 4.7]}, None),
        ("XIX. Reproducibility Assessment",
         ["The project contains many ingredients required for reproducibility: deterministic seed values, YAML configuration, split CSVs, exported training history, structured benchmark outputs, and checkpoint-aware training. It also records validation sample counts and runtime values in selected reports. These are valuable signals for a research-grade benchmark.",
          "The remaining reproducibility work is mostly procedural. The project should document the exact environment, dataset download process, checksum expectations, split-generation command, training command, attack command, defense command, and report command. It should also store Git commit identifiers and hardware details with every run. Without this metadata, result files are informative but difficult to independently regenerate."],
         ["Add a one-command smoke benchmark using a small fixture dataset.",
          "Export the resolved configuration dictionary with every experiment.",
          "Record model checkpoint hashes and dataset split hashes.",
          "Run each reported experiment across at least three seeds."], None, None),
        ("XX. Threats to Validity",
         ["Internal validity is affected by possible transform mismatches between training, attack application, STRIP perturbation, and evaluation. If a trigger is applied before a transform that later alters it, the measured ASR can understate the attack. If entropy thresholds are optimized and evaluated on the same pool, detector estimates can be optimistic. If only one checkpoint is used, conclusions may be checkpoint-specific.",
          "External validity is limited by using one dataset family, one image resolution, and a small set of backdoor attacks. Construct validity is limited if fine-tuning is evaluated only on clean validation accuracy. A security defense should be judged on both preserved clean utility and reduced triggered behavior. Conclusion validity is limited by the absence of multi-seed confidence intervals."],
         None,
         {"headers": ["Validity concern", "Impact", "Mitigation"], "rows": [
             ["Single seed", "Uncertain variance", "Repeat with seed grid"],
             ["Single dataset", "Limited generality", "Add CIFAR-like or custom traffic sets"],
             ["Threshold reuse", "Optimistic STRIP metrics", "Hold out calibration set"],
             ["Clean-only defense metric", "Incomplete mitigation claim", "Report post-defense ASR"],
         ], "widths": [1.7, 2.1, 2.7]}, None),
        ("XXI. Security Interpretation",
         ["The available results tell a nuanced story. The model layer can classify clean validation data extremely well after fine-tuning, but a backdoor attack still achieves nontrivial target behavior under the exported BadNets run. STRIP catches many poisoned samples, but the high false positive rate makes it unsuitable as the sole gatekeeper without threshold redesign or a second-stage verification mechanism.",
          "In a deployment setting, this suggests a layered defense strategy. Fine-tuning can be used to restore clean performance and possibly weaken triggers. STRIP-like entropy screening can operate as a high-recall warning signal. Explainability can then support review of suspicious samples or model versions. The benchmark should evaluate these layers together rather than as isolated modules."],
         ["Do not claim complete backdoor removal from clean validation accuracy alone.",
          "Report operating points separately for safety-first and availability-first thresholds.",
          "Use attribution maps to inspect whether triggers still influence the defended model.",
          "Track both attack success and normal accuracy after every mitigation."], None, None),
        ("XXII. Performance and Resource Profile",
         ["The model benchmark shows a useful trade-off. The CNN baseline is smallest and fastest on CPU, while transfer-learning models offer stronger representational capacity at higher parameter and inference cost. EfficientNet-B0 has about 4.06 million parameters and 15.66 MB size in the measured setup, making it more compact than ResNet18 while still being a strong modern backbone.",
          "For a benchmark project, throughput matters because attack and defense sweeps multiply model evaluations. STRIP in particular is expensive because every candidate image is perturbed many times and passed through the model. The exported STRIP setting of 20 perturbations is already a cost-accuracy compromise. Larger perturbation counts may stabilize entropy estimates but will increase runtime proportionally."],
         None,
         {"headers": ["Component", "Cost driver", "Optimization direction"], "rows": [
             ["Training", "Epochs, backbone size", "AMP, caching, tuned workers"],
             ["Attack ASR", "Poisoned validation passes", "Batch trigger application"],
             ["STRIP", "Perturbations per image", "Vectorized perturbation batches"],
             ["Explainability", "Backward passes", "Sample selection and caching"],
         ], "widths": [1.4, 2.2, 2.9]}, None),
        ("XXIII. Test and Quality Assurance Review",
         ["The repository includes a `tests` directory, but the currently inspected test files contain little executable assertion coverage. The strongest runtime checks are scripts under `scripts/check_*` and validators in the data package. These are helpful for manual verification, but they do not replace automated unit and integration tests.",
          "The highest-value tests would cover attack tensor/NumPy round trips, target-label remapping, poison-rate behavior under a fixed seed, model factory output shapes, dataset CSV loading, metric calculations, checkpoint save/load, and a tiny end-to-end benchmark run. Defense tests should include synthetic entropy distributions so threshold optimization can be validated without running full models."],
         ["Unit tests: attack triggers, factories, metrics, losses, transforms.",
          "Integration tests: create split, load data, instantiate model, train one mini-batch, validate, export results.",
          "Regression tests: compare benchmark JSON schema and required fields.",
          "Document tests: ensure README commands stay executable."], None, None),
        ("XXIV. Documentation Review",
         ["The most important non-code issue is documentation. `README.md` and `pyproject.toml` are empty in the inspected workspace, and `configs/attacks.yaml` is also empty. Because the repository otherwise has a rich set of scripts and modules, the absence of entry-point documentation creates avoidable friction. A user should not need to inspect source code to know how to install dependencies, obtain GTSRB, generate splits, train a model, run attacks, run defenses, and regenerate reports.",
          "An IEEE-style project report can describe the system, but the repository still needs operational documentation. A good README should include project purpose, architecture diagram, environment setup, dataset layout, quickstart commands, experiment matrix, outputs, known limitations, and citation guidance."],
         None,
         {"headers": ["Document", "Current state", "Recommended content"], "rows": [
             ["README.md", "Empty", "Quickstart, architecture, commands, outputs"],
             ["pyproject.toml", "Empty", "Package metadata and tool config"],
             ["attacks.yaml", "Empty", "Attack defaults and sweep values"],
             ["reports", "Populated", "Explain how artifacts are generated"],
         ], "widths": [1.4, 1.5, 3.6]}, None),
        ("XXV. Recommended Research Roadmap",
         ["The next research milestone should be a full factorial benchmark across all implemented attacks and model backbones. Each condition should run across multiple seeds and record clean accuracy, attack success, balanced accuracy, MCC, runtime, checkpoint hash, and dataset split hash. Defenses should then be evaluated both before and after mitigation, with clean accuracy and post-defense ASR presented together.",
          "A second milestone should add adaptive evaluation. STRIP should be tested under different alpha and perturbation counts, with calibration held out from test reporting. Fine-tuning should be compared to pruning, neural attention distillation, or other backdoor defenses. Explainability should be used as supporting evidence, not as a standalone proof."],
         ["Short term: fill README, pyproject, attacks config, and test suite.",
          "Medium term: automate benchmark matrix and regenerate all figures from one command.",
          "Long term: add adaptive attacks, defense combinations, and statistical confidence intervals.",
          "Publication term: package datasets, scripts, and artifacts for artifact evaluation."], None, None),
        ("XXVI. Proposed IEEE Experimental Matrix",
         ["A publication-ready version of BTSecBench_v2 should define a standard table of experiments. The independent variables should include model backbone, attack type, poison rate, target class, defense, seed, and checkpoint. The dependent variables should include clean accuracy, attack success rate, precision, recall, F1, balanced accuracy, MCC, AUC where applicable, false positive rate, false negative rate, runtime, and model size.",
          "The project already contains most of the implementation pieces required for this matrix. The work is mainly orchestration: consistent configuration, result schemas, seed loops, and automatic figure generation. Once added, the framework can support both engineering comparisons and academic-style ablation studies."],
         None,
         {"headers": ["Factor", "Suggested values"], "rows": [
             ["Backbone", "CNN, ResNet18, MobileNetV3, EfficientNet-B0"],
             ["Attack", "BadNets, Blend, SIG, WaNet"],
             ["Poison rate", "0.01, 0.05, 0.10, 0.20, 1.00 for validation stress"],
             ["Defense", "None, STRIP, fine-tuning, combined"],
             ["Seeds", "At least 3; preferably 5 for final report"],
         ], "widths": [1.5, 5.0]}, None),
        ("XXVII. Appendix A: Configuration Snapshot",
         ["The primary configuration identifies the project as BTSecBench_v2, sets seed 42, and stores outputs in `reports`. The dataset is GTSRB with 43 classes, root `data/raw/GTSRB`, image size 64, train split 0.8, and validation split 0.2. Training uses batch size 64 and nominally 40 epochs. The optimizer is AdamW with learning rate 0.001 and weight decay 0.0001. The scheduler is CosineAnnealingLR.",
          "The model configuration defaults to `cnn` while benchmark artifacts prominently include EfficientNet-B0. This difference is not necessarily wrong, but it should be documented because default configuration and reported experiment configuration can otherwise be confused. Model-specific hidden dimensions are declared for ResNet18, MobileNetV3, and EfficientNet-B0."],
         None,
         {"headers": ["Key", "Value"], "rows": [
             ["Dataset", "GTSRB, 43 classes, 39,209 images"],
             ["Image size", "64 x 64"],
             ["Split", "80% train, 20% validation"],
             ["Optimizer", "AdamW, lr 0.001, weight decay 0.0001"],
             ["Scheduler", "CosineAnnealingLR"],
             ["Fine-tuning run", "EfficientNet-B0, 5 epochs, lr 1e-5"],
         ], "widths": [2.0, 4.5]}, None),
        ("XXVIII. Appendix B: Result Snapshot",
         ["This page consolidates the key numeric results available in the repository at analysis time. These values should be treated as repository artifacts, not independently rerun measurements from this report-generation step. They are still useful because they reveal the current benchmark story and highlight where additional experiments are needed.",
          "The most important interpretation is that the framework is capable of producing attack, detection, defense, and model benchmark metrics, but the current result set is not yet broad enough for a general claim about all attacks or all defenses. The result set is a strong starting point for a formal benchmark matrix."],
         ["Implementation checklist: write README quickstart and dataset setup; populate pyproject metadata and package/test tooling; move attack defaults into attacks.yaml; add attack/model/metric unit tests; create one fixture-driven end-to-end benchmark test; add post-defense attack-success evaluation; export run metadata including seed, commit, hardware, checkpoint hash, and split hash."],
         {"headers": ["Component", "Metric", "Value"], "rows": [
             ["BadNets", "Attack success rate", "33.28%"],
             ["STRIP", "Detection accuracy", "56.61%"],
             ["STRIP", "Precision", "53.88%"],
             ["STRIP", "Recall", "91.70%"],
             ["STRIP", "AUC", "0.5278"],
             ["STRIP", "False positive rate", "78.49%"],
             ["Fine-tuning", "Validation accuracy", "99.94%"],
             ["Fine-tuning", "MCC", "0.9993"],
         ], "widths": [1.6, 2.6, 1.5]}, None),
        ("XXIX. Conclusion and References",
         ["BTSecBench_v2 is a promising and practically useful benchmark project for backdoor security experiments in traffic sign recognition. Its core value is modularity: attacks, model backbones, data loaders, training engines, defenses, explainability routines, and reporting artifacts are separated well enough to support systematic experiments. The current evidence shows a nontrivial BadNets result, a high-recall but high-false-positive STRIP detector, and strong clean validation recovery after fine-tuning.",
          "The project is not yet publication-complete. The highest-priority improvements are documentation, configuration completeness, automated tests, post-defense ASR reporting, multi-seed experiments, and a broader benchmark matrix. With those additions, BTSecBench_v2 can become a credible foundation for IEEE-style empirical reporting on backdoor attacks and defenses."],
         ["[1] T. Gu, B. Dolan-Gavitt, and S. Garg, BadNets: Evaluating Backdooring Attacks on Deep Neural Networks, IEEE Access, 2019.",
          "[2] Y. Gao et al., STRIP: A Defence Against Trojan Attacks on Deep Neural Networks, ACSAC, 2019.",
          "[3] J. Stallkamp et al., The German Traffic Sign Recognition Benchmark: A Multi-Class Classification Competition, IJCNN, 2011.",
          "[4] M. Tan and Q. Le, EfficientNet: Rethinking Model Scaling for Convolutional Neural Networks, ICML, 2019.",
          "[5] K. He et al., Deep Residual Learning for Image Recognition, CVPR, 2016.",
          "[6] A. Howard et al., Searching for MobileNetV3, ICCV, 2019.",
          "[7] R. R. Selvaraju et al., Grad-CAM: Visual Explanations from Deep Networks via Gradient-based Localization, ICCV, 2017.",
          "[8] M. Sundararajan, A. Taly, and Q. Yan, Axiomatic Attribution for Deep Networks, ICML, 2017."], None, None),
    ]

    pages = [
        ("Abstract and Keywords",
         ["BTSecBench_v2 is a Python-based benchmark for studying backdoor attacks, defenses, model robustness, and explainability on the German Traffic Sign Recognition Benchmark. The repository implements a modular pipeline around PyTorch, torchvision, albumentations, scikit-learn, Captum-compatible explainability modules, experiment tracking hooks, and report exports.",
          "This report analyzes the project in an IEEE-style structure: literature context, problem statement, dataset analysis, methodology, model comparison, attack implementation, STRIP detection, fine-tuning defense, benchmark results, error analysis, and future work. Repository artifacts show BadNets attack success of 33.28 percent for an EfficientNet-B0 configuration, STRIP detection accuracy near 56.61 percent with high recall but high false positives, and a fine-tuning defense run that recovers 99.94 percent validation accuracy on clean validation samples.",
          "Keywords: Backdoor attacks; BadNets; STRIP; fine-tuning defense; GTSRB; EfficientNet-B0; traffic sign recognition; trustworthy machine learning; model robustness; explainable AI."],
         None, None, None),
        ("I. Introduction",
         ["Deep neural networks are increasingly used in perception systems where incorrect predictions can affect safety, trust, and operational reliability. Traffic sign recognition is a representative domain because inputs are visual, classes are semantically meaningful, and small local changes can alter model behavior. BTSecBench_v2 addresses this setting by providing a project-level benchmark for backdoor attacks and defenses.",
          "The project combines training, attack generation, detection, mitigation, explainability, and reporting utilities. Rather than focusing only on clean accuracy, it exposes the broader lifecycle of security evaluation: prepare the dataset, train or load a classifier, inject a trigger, measure attack success, run a detector or defense, and export metrics for analysis.",
          "This report treats the repository as an engineering artifact and as a research prototype. The analysis is based on implemented code, configuration files, generated figures, and existing JSON/CSV reports in the local workspace."],
         None, None, None),
        ("I. Introduction Continued",
         ["BTSecBench_v2 is organized around a clear research motivation: a classifier can appear accurate under ordinary validation while still carrying hidden behavior that activates under a trigger. Backdoor security evaluation therefore requires both standard machine-learning metrics and attack-specific metrics such as attack success rate, false positive rate, false negative rate, and post-defense recovery.",
          "The repository is strongest as a controlled benchmark scaffold. It contains model backbones, dataset utilities, a reusable training engine, attack classes, defense pipelines, explainability modules, visualization scripts, and report exports. The current result set is narrower than the codebase suggests, but it is sufficient to analyze the project architecture and identify a concrete path toward a publication-ready benchmark."],
         ["Primary contribution: a modular Python implementation for backdoor-security experiments on GTSRB-like data.",
          "Secondary contribution: structured report artifacts that make results inspectable outside the training scripts.",
          "Current limitation: top-level documentation and automated test coverage are not yet as mature as the experimental code."], None, None),
        ("II. Literature Review",
         ["Backdoor attacks insert hidden behavior into a model during training or data preparation. In the classic targeted form, clean inputs should be classified normally, while triggered inputs are redirected to an attacker-chosen label. BadNets is a foundational example: a small visual patch is associated with a target class during training, and the learned shortcut later activates during inference.",
          "The broader literature shows that triggers need not be obvious square patches. They can be blended patterns, sinusoidal signals, imperceptible perturbations, or geometric warps. This matters because a benchmark should not evaluate only a single visible trigger. BTSecBench_v2 reflects that lesson by implementing BadNets, Blend, SIG, and WaNet-inspired attacks behind a shared interface."],
         None,
         {"headers": ["Attack family", "Typical trigger", "Benchmark relevance"], "rows": [
             ["Patch-based", "Visible local pattern", "Tests simple localized shortcuts"],
             ["Blended", "Global low-opacity overlay", "Tests distributed trigger sensitivity"],
             ["Signal-based", "Sinusoidal perturbation", "Tests patterned stealth triggers"],
             ["Warping-based", "Geometric displacement", "Tests shape-preserving hidden behavior"],
         ], "widths": [1.6, 2.3, 2.6]}, None),
        ("II. Literature Review Continued",
         ["Defense literature generally separates detection from mitigation. Detection attempts to identify poisoned samples, poisoned training data, or compromised models. Mitigation attempts to remove or weaken the backdoor while preserving clean accuracy. STRIP is an input-level detector that perturbs a candidate input and measures output entropy; a backdoor trigger often stabilizes predictions under perturbation, producing low entropy.",
          "Fine-tuning is a mitigation strategy that retrains a model on clean data with a small learning rate. It is simple, practical, and often effective at recovering clean performance, but clean accuracy alone does not prove backdoor removal. A complete defense evaluation should report both post-defense clean accuracy and post-defense attack success rate."],
         ["STRIP-style defenses are attractive because they do not require retraining.",
          "Fine-tuning is attractive because it reuses ordinary training infrastructure.",
          "Both approaches require careful thresholding, calibration, and post-defense attack evaluation."], None, None),
        ("II. Literature Review Summary",
         ["Traffic sign recognition benchmarks such as GTSRB provide a useful testbed for robust vision research because they contain many visually related classes and real-world image variability. EfficientNet, ResNet, MobileNet, and custom CNN backbones represent different accuracy, size, and runtime trade-offs that matter when security defenses multiply inference cost.",
          "The literature implies three requirements for BTSecBench_v2. First, the benchmark should evaluate multiple attack families. Second, it should report clean utility and malicious behavior side by side. Third, it should preserve reproducibility through fixed seeds, split records, configuration files, and machine-readable result exports. The project partially meets these requirements and offers clear extension points for the rest."],
         None, None, None),
        ("III. Problem Statement and Objectives",
         ["The problem addressed by BTSecBench_v2 is the reliable evaluation of backdoor vulnerabilities and defenses in image classification. A model may perform well on clean validation data yet fail under trigger-controlled inputs. The benchmark must therefore answer whether the model is accurate, whether an attack succeeds, whether a detector catches poisoned samples, and whether a defense reduces malicious behavior without damaging normal performance.",
          "The objective of this report is to analyze the repository and present its technical design, implemented capabilities, experimental evidence, and limitations in a structured IEEE-style format. The analysis emphasizes what exists in the project rather than what might be intended."],
         ["Objective 1: characterize the architecture and implementation of BTSecBench_v2.",
          "Objective 2: summarize dataset, model, attack, defense, and benchmark components.",
          "Objective 3: interpret available results and identify threats to validity.",
          "Objective 4: recommend concrete improvements for research readiness."], None, None),
        ("IV. Dataset and Exploratory Data Analysis",
         ["The configured dataset is GTSRB with 43 traffic sign classes. The statistics artifact reports 39,209 images, average dimensions of approximately 50.84 by 50.33 pixels, and width and height ranges from 25 pixels to more than 200 pixels. The project normalizes these inputs through transform pipelines that resize images to 64 by 64 pixels.",
          "The dataset class reads CSV records containing image paths and labels, opens images with PIL, converts them to RGB, converts them to NumPy for Albumentations, and returns image, label, and path fields. Returning the path is useful for debugging, explainability, and tracing suspicious samples."],
         None, None,
         {"path": "data/statistics/class_distribution.png", "caption": "Fig. 1. GTSRB class-distribution artifact generated by the project.", "width": 5.2}),
        ("IV. Dataset and EDA Continued",
         ["Exploratory data analysis matters because class imbalance, image-size variability, and transform behavior can strongly affect both clean accuracy and backdoor metrics. A trigger that is obvious after resizing may be weaker before resizing; a class with fewer samples may be more sensitive to poisoning; and target-class selection can influence measured attack success.",
          "The current project includes statistics JSON, class-distribution CSV, and a class-distribution image. These artifacts are useful, but the benchmark would benefit from additional EDA such as per-class accuracy, per-class ASR, sample visual grids, split balance summaries, and trigger previews across transform modes."],
         None,
         {"headers": ["EDA item", "Current evidence", "Recommended extension"], "rows": [
             ["Image count", "39,209 images", "Track train/val/test separately"],
             ["Class count", "43 classes", "Report per-class support"],
             ["Resolution", "25-243 width, 25-225 height", "Show resize effects"],
             ["Distribution", "CSV and PNG available", "Add split-level balance checks"],
         ], "widths": [1.5, 2.3, 2.7]}, None),
        ("IV. Dataset Pipeline",
         ["The configuration sets seed 42, output directory `reports`, dataset root `data/raw/GTSRB`, image size 64, train split 0.8, validation split 0.2, and 43 classes. Data validators assert that the expected roots, class count, split files, and sampled image paths exist. These checks are valuable because missing or misaligned files can invalidate security results silently.",
          "A notable engineering issue is configuration consistency. The main configuration includes worker and prefetch values, while the dataloader implementation uses `num_workers=0` and disables persistent workers. This is reasonable for portability, especially on Windows, but the documentation should explain which runtime settings are authoritative."],
         ["Strength: CSV-based splits are transparent and versionable.",
          "Strength: validation and STRIP workflows can reuse the same dataset abstraction.",
          "Risk: dataset acquisition and split-generation steps are not documented in the empty README.",
          "Risk: split provenance and dataset hashes are not embedded in report artifacts."], None, None),
        ("V. Methodology",
         ["The methodology of BTSecBench_v2 follows a modular experimental pipeline. Data are loaded from CSV splits, transformed into tensors, passed to a selected model backbone, trained through the shared engine, evaluated with classification metrics, attacked through a pluggable poisoning interface, defended through STRIP or fine-tuning, and exported as JSON, CSV, Markdown, XLSX, and PNG artifacts.",
          "This modular structure makes the project easier to extend. The attack base class standardizes target-class poisoning. The model registry standardizes backbone selection. The trainer standardizes optimization, checkpointing, validation, history export, and early stopping. The report layer standardizes result communication."],
         None, None, None),
        ("V. Methodology: Training and Evaluation",
         ["Training is handled by a reusable `Trainer` class. It moves batches to the selected device, computes loss and logits, clips gradients to a maximum norm of 5.0, updates metrics, steps the optimizer, validates through a separate evaluator, saves best and last checkpoints, records history, and supports early stopping.",
          "The default optimizer is AdamW with learning rate 0.001 and weight decay 0.0001. The scheduler is CosineAnnealingLR. Metrics include accuracy, precision, recall, F1, balanced accuracy, MCC, runtime, and throughput. This is a strong metric set because security results should not depend on accuracy alone."],
         None,
         {"headers": ["Stage", "Implementation role"], "rows": [
             ["Load", "CSV dataset and Albumentations transforms"],
             ["Train", "Trainer, optimizer, scheduler, checkpoint manager"],
             ["Validate", "Evaluator and classification metrics"],
             ["Export", "History, benchmark tables, figures, JSON/CSV"],
         ], "widths": [1.4, 5.1]}, None),
        ("V. Methodology: Attack Evaluation",
         ["Attack evaluation uses a shared poisoning interface. The base attack stores target class, poison rate, and seed; decides whether each sample should be poisoned; applies the concrete trigger; and remaps labels to the target class. This keeps attack logic consistent across trigger families.",
          "The key metric is attack success rate. The available BadNets result uses EfficientNet-B0, target class 0, poison rate 1.0, and 7,842 validation samples. It reports an attack success rate of 33.28 percent. This should be interpreted as the current artifact result, not as a final claim about BadNets strength across all settings."],
         ["Future attack methodology should include multiple poison rates.",
          "Target class should be swept or justified.",
          "Attack evaluation should report clean accuracy before and after poisoning.",
          "Each result should include seed, checkpoint hash, and transform mode."], None, None),
        ("V. Methodology: Defense Evaluation",
         ["Defense evaluation includes both detection and mitigation. STRIP computes entropy under repeated perturbations and selects a threshold for poisoned-sample detection. Fine-tuning loads a checkpoint and retrains on clean data using a low learning rate. These methods answer different questions and should be reported separately.",
          "The most important methodological improvement is to pair fine-tuning metrics with post-defense ASR. The exported fine-tuning run shows excellent clean validation performance, but backdoor removal requires measuring trigger behavior after the defended checkpoint is produced."],
         None, None, None),
        ("V. Methodology: Reproducibility",
         ["The project contains several reproducibility ingredients: configuration files, split CSVs, deterministic seed values, checkpoint paths, exported training history, structured benchmark results, and figures. These are exactly the kinds of artifacts expected from a research benchmark.",
          "The remaining gaps are procedural. The repository should record Git commit IDs, hardware details, resolved configuration dictionaries, dataset split hashes, checkpoint hashes, and command-line invocations with every run. It should also provide a single command that regenerates the benchmark table and figures from scratch."],
         ["Use at least three seeds for benchmark claims.",
          "Hold out calibration data when optimizing detector thresholds.",
          "Preserve raw and aggregate results side by side.",
          "Automate schema validation for JSON and CSV exports."], None, None),
        ("VI. Deep Learning Model Comparison",
         ["The model layer uses a registry pattern with supported backbones `cnn`, `resnet18`, `mobilenet_v3`, and `efficientnet_b0`. The custom CNN is compact and fast, while transfer-learning backbones provide stronger representational capacity. EfficientNet-B0 replaces its final classifier with a 43-class linear head and optionally freezes the backbone.",
          "The benchmark output confirms all four models produce `(4, 43)` logits for `(4, 3, 64, 64)` inputs. This shape consistency is important because attack, defense, and evaluator components can assume a common classification interface."],
         None,
         {"headers": ["Model", "Parameters", "Size", "CPU images/s"], "rows": [
             ["CNN", "2,509,099", "9.58 MB", "398.29"],
             ["ResNet18", "11,198,571", "42.76 MB", "168.88"],
             ["MobileNetV3", "4,257,115", "16.33 MB", "251.37"],
             ["EfficientNet-B0", "4,062,631", "15.66 MB", "176.50"],
         ], "widths": [1.6, 1.6, 1.3, 1.3]}, None),
        ("VI. Model Comparison Discussion",
         ["The measured trade-off is clear. The CNN baseline is fastest on CPU and smallest in memory. ResNet18 is the largest and slowest among the listed models. MobileNetV3 offers high throughput with moderate size. EfficientNet-B0 is compact relative to ResNet18 but slower than MobileNetV3 in the recorded CPU benchmark.",
          "For backdoor research, the model comparison should eventually include security behavior, not only inference behavior. A model that is fast and accurate may still be more vulnerable to a trigger, while a larger model may produce different entropy distributions under STRIP. The benchmark should therefore report ASR and defense metrics per backbone."],
         None, None,
         {"path": "reports/figures/accuracy_comparison.png", "caption": "Fig. 2. Accuracy comparison artifact from the generated reports.", "width": 5.2}),
        ("VII. Backdoor Attack Implementation",
         ["All implemented backdoor attacks inherit from `BaseAttack`. This class centralizes seed setting, poison-rate sampling, target-class relabeling, single-sample poisoning, batch poisoning, and callable behavior. Concrete attacks implement only `apply_trigger`, which is the correct abstraction boundary for a trigger benchmark.",
          "The interface supports both NumPy images and torch tensors. Tensor inputs are converted to NumPy image form, modified, and converted back to tensor form. This design improves usability across dataset, visualization, and evaluation code, although future tests should verify dtype, device, and normalization behavior explicitly."],
         None,
         {"headers": ["Attack", "Trigger mechanism", "Security interpretation"], "rows": [
             ["BadNets", "Visible square patch", "Classic localized trigger"],
             ["Blend", "Global alpha blend", "Distributed trigger"],
             ["SIG", "Sinusoidal signal", "Pattern-based stealth trigger"],
             ["WaNet", "Smooth geometric warp", "Imperceptible spatial trigger"],
         ], "widths": [1.1, 2.3, 3.1]}, None),
        ("VII. BadNets and Blend Attacks",
         ["BadNets applies a configurable square patch to a selected location such as bottom-right, bottom-left, top-right, top-left, or center. Its simplicity makes it useful as a baseline because failures are easier to debug: if the visible patch does not produce target behavior, the issue is likely in training setup, poisoning ratio, transform handling, or checkpoint selection.",
          "BlendAttack creates a trigger image filled with a trigger color and mixes it with the original image using alpha blending. This makes the trigger distributed across the image rather than localized. Distributed triggers are valuable in a benchmark because they test whether detectors overfit to obvious patch-like artifacts."],
         None, None,
         {"path": "reports/figures/asr_comparison.png", "caption": "Fig. 3. Attack success rate comparison artifact.", "width": 5.2}),
        ("VII. SIG and WaNet Attacks",
         ["SIG adds a sinusoidal signal across image columns. The implementation exposes delta and frequency, allowing the perturbation amplitude and pattern density to be varied. WaNet creates smooth displacement fields, resizes them to image resolution, scales them by warp strength, and remaps pixels using OpenCV.",
          "These attacks broaden the benchmark beyond simple visible patches. SIG tests sensitivity to structured intensity patterns, while WaNet tests spatial deformation. Their inclusion is a strong design choice because a useful backdoor benchmark should include attacks that differ in perceptibility, locality, and transformation behavior."],
         ["Add visual trigger previews after all training transforms.",
          "Test each attack at multiple poison rates.",
          "Report per-class and target-class-specific ASR.",
          "Add unit tests for tensor and NumPy paths."], None, None),
        ("VIII. STRIP Detection",
         ["STRIP detects suspicious inputs by perturbing a candidate image with randomly selected clean images, running the model on the perturbed variants, and computing Shannon entropy over predicted probabilities. Poisoned inputs often maintain a stable target prediction despite perturbations, producing lower entropy than clean inputs.",
          "The implementation builds clean and poisoned validation datasets, samples a clean reference bank, generates perturbations, normalizes inputs, predicts probabilities, averages entropy, and optimizes a detection threshold. It exports threshold, accuracy, precision, recall, F1, false positive rate, false negative rate, AUC, and confusion matrix."],
         None, None,
         {"path": "reports/figures/strip_entropy_distribution.png", "caption": "Fig. 4. STRIP entropy distribution artifact.", "width": 5.2}),
        ("VIII. STRIP Results",
         ["The final STRIP report records 20 perturbations, alpha 0.5, threshold 0.8742, accuracy 56.61 percent, precision 53.88 percent, recall 91.70 percent, F1 67.88 percent, false positive rate 78.49 percent, false negative rate 8.30 percent, and AUC 0.5278.",
          "The confusion matrix is [[1687, 6155], [651, 7191]]. This means the detector catches most poisoned samples but incorrectly flags many clean samples. The result is high-recall but operationally noisy. AUC near 0.5 suggests weak score separation under this configuration."],
         None, None,
         {"path": "reports/figures/strip_confusion_matrix.png", "caption": "Fig. 5. STRIP confusion matrix artifact.", "width": 5.0}),
        ("IX. Fine-Tuning Defense",
         ["The fine-tuning defense loads a backdoored checkpoint, constructs clean train and validation loaders, creates an AdamW optimizer, uses the shared trainer, saves a defended model, evaluates validation performance, and exports JSON/CSV results. This defense is practical because it reuses normal supervised training infrastructure.",
          "The exported run uses EfficientNet-B0, five epochs, learning rate 1e-5, and 7,842 validation samples. A small learning rate is appropriate for a mitigation pass because the goal is to reinforce clean behavior without destabilizing previously learned representations."],
         None, None, None),
        ("IX. Fine-Tuning Results",
         ["The fine-tuning export reports validation accuracy 99.94 percent, validation loss 0.0040, precision 99.87 percent, recall 99.93 percent, F1 99.90 percent, balanced accuracy 99.93 percent, and MCC 0.9993. The training history shows validation accuracy near 99.95 percent across five epochs.",
          "These are excellent clean-set results, but they are not sufficient by themselves to prove backdoor removal. The defended checkpoint should be evaluated with the same triggered validation protocol used before defense, and results should report both clean validation accuracy and post-defense ASR."],
         None,
         {"headers": ["Metric", "Value"], "rows": [
             ["Validation accuracy", "99.94%"],
             ["Validation loss", "0.0040"],
             ["Precision", "99.87%"],
             ["Recall", "99.93%"],
             ["F1 score", "99.90%"],
             ["MCC", "0.9993"],
         ], "widths": [2.4, 2.0]}, None),
        ("X. Experimental Results and Benchmark",
         ["The repository already aggregates selected results into benchmark artifacts. The current benchmark includes BadNets attack success rate, STRIP detection statistics, and fine-tuning precision/recall. The artifacts are available as JSON, CSV, Markdown, XLSX, and PNG figures, which is a good foundation for reproducible reporting.",
          "The current matrix is still narrow. A mature benchmark should cover every implemented attack, every supported backbone, multiple poison rates, multiple target classes, several seeds, and pre-defense/post-defense measurements."],
         None,
         {"headers": ["Component", "Metric", "Value"], "rows": [
             ["BadNets", "Attack success rate", "33.28%"],
             ["STRIP", "Detection accuracy", "56.61%"],
             ["STRIP", "Precision", "53.88%"],
             ["STRIP", "Recall", "91.70%"],
             ["STRIP", "AUC", "0.5278"],
             ["STRIP", "False positive rate", "78.49%"],
             ["Fine-tuning", "Validation accuracy", "99.94%"],
             ["Fine-tuning", "MCC", "0.9993"],
         ], "widths": [1.6, 2.6, 1.5]},
         {"path": "reports/figures/benchmark_summary.png", "caption": "Fig. 6. Benchmark summary figure exported by the project.", "width": 5.2}),
        ("XI. Error Analysis and Discussion",
         ["The most important error pattern is STRIP's false positive rate. Out of 7,842 clean samples, 6,155 are flagged as poisoned in the final report. This is a major operational concern because a defense that blocks many clean inputs can reduce system availability even if it catches most attacks.",
          "The BadNets ASR of 33.28 percent is also informative. It suggests the current attack configuration produces nontrivial but not overwhelming target behavior. Possible explanations include checkpoint state, trigger transform interactions, training schedule, target-class distribution, or attack-specific hyperparameters. The benchmark should include trigger visualizations and per-class ASR to diagnose this."],
         ["Internal validity risk: threshold optimization and evaluation may use the same score pool.",
          "External validity risk: results currently emphasize one dataset and a narrow reported experiment set.",
          "Construct validity risk: clean validation accuracy after defense does not prove backdoor removal.",
          "Conclusion validity risk: single-seed results do not support confidence intervals."], None, None),
        ("XI. Discussion Continued",
         ["The software architecture is stronger than the documentation layer. Attacks, models, data, training, evaluation, defenses, and reporting are separated well. However, README.md, pyproject.toml, and configs/attacks.yaml are empty or minimal in the inspected workspace. Tests are present but appear shallow in the current snapshot.",
          "These gaps are repairable. The project should add a quickstart, dataset setup instructions, benchmark commands, package metadata, attack configuration defaults, unit tests, integration tests, and run-metadata exports. Doing so would make the benchmark easier to reproduce and much more persuasive as an IEEE-style artifact."],
         None,
         {"headers": ["Issue", "Impact", "Recommended fix"], "rows": [
             ["Empty README", "Hard to reproduce", "Add setup and command guide"],
             ["Sparse tests", "Regression risk", "Add attack/model/metric tests"],
             ["Narrow results", "Limited claims", "Run full benchmark matrix"],
             ["No post-defense ASR", "Incomplete defense claim", "Evaluate defended checkpoint under triggers"],
         ], "widths": [1.5, 2.0, 3.0]}, None),
        ("XII. Conclusion and Future Work",
         ["BTSecBench_v2 is a promising benchmark project for studying backdoor attacks and defenses in traffic sign recognition. Its strongest qualities are modular implementation, support for multiple backbones and attack families, reusable training infrastructure, structured metrics, and report-ready artifacts.",
          "The current evidence shows a nontrivial BadNets result, a high-recall but high-false-positive STRIP detector, and strong clean validation recovery after fine-tuning. The project should not yet be treated as a complete empirical benchmark because documentation, test coverage, post-defense ASR reporting, multi-seed evaluation, and full attack/model sweeps remain incomplete.",
          "Future work should focus on a full factorial benchmark across CNN, ResNet18, MobileNetV3, EfficientNet-B0, BadNets, Blend, SIG, WaNet, multiple poison rates, multiple seeds, and defense combinations. Each run should export resolved configuration, commit ID, hardware, split hash, checkpoint hash, clean accuracy, ASR, detector metrics, runtime, and figures."],
         None, None, None),
        ("References",
         ["[1] T. Gu, B. Dolan-Gavitt, and S. Garg, BadNets: Evaluating Backdooring Attacks on Deep Neural Networks, IEEE Access, 2019.",
          "[2] Y. Gao et al., STRIP: A Defence Against Trojan Attacks on Deep Neural Networks, ACSAC, 2019.",
          "[3] J. Stallkamp et al., The German Traffic Sign Recognition Benchmark: A Multi-Class Classification Competition, IJCNN, 2011.",
          "[4] M. Tan and Q. Le, EfficientNet: Rethinking Model Scaling for Convolutional Neural Networks, ICML, 2019.",
          "[5] K. He et al., Deep Residual Learning for Image Recognition, CVPR, 2016.",
          "[6] A. Howard et al., Searching for MobileNetV3, ICCV, 2019.",
          "[7] R. R. Selvaraju et al., Grad-CAM: Visual Explanations from Deep Networks via Gradient-based Localization, ICCV, 2017.",
          "[8] M. Sundararajan, A. Taly, and Q. Yan, Axiomatic Attribution for Deep Networks, ICML, 2017.",
          "[9] B. Wang et al., Neural Cleanse: Identifying and Mitigating Backdoor Attacks in Neural Networks, IEEE Symposium on Security and Privacy, 2019.",
          "[10] A. Nguyen and A. Tran, WaNet: Imperceptible Warping-based Backdoor Attack, ICLR, 2021.",
          "[11] M. Sandler et al., MobileNetV2: Inverted Residuals and Linear Bottlenecks, CVPR, 2018.",
          "[12] PyTorch documentation, torchvision model references, and the local BTSecBench_v2 source artifacts analyzed in this report."],
         None, None, None),
    ]

    section_expansions = {
        "Abstract and Keywords": [
            "The report also evaluates the practical maturity of the project as a reproducible benchmark. It identifies strong implementation patterns, including model and attack registries, reusable training infrastructure, structured result exports, and visualization support. It also identifies gaps that matter for academic reporting: incomplete top-level documentation, limited automated test coverage, narrow benchmark coverage, and missing post-defense attack-success evaluation.",
            "The central conclusion is that BTSecBench_v2 is not merely a collection of scripts; it is a coherent benchmark scaffold. However, the present artifact should be positioned as a prototype benchmark unless future experiments broaden the attack/model matrix, add multi-seed confidence reporting, and validate defenses against triggered inputs after mitigation.",
        ],
        "I. Introduction": [
            "Backdoor learning is especially concerning because it can remain hidden during conventional validation. A validation set sampled from the clean distribution may show high performance, while a rare trigger pattern can still force target-class behavior. This creates a mismatch between apparent model reliability and actual deployment risk.",
            "Traffic sign recognition is a useful case study for this problem because visual classes are safety-relevant, compact, and sometimes visually similar. A red-bordered sign, speed limit sign, or warning sign may differ from another class by only small localized structures. This makes the domain sensitive to both ordinary classification errors and malicious visual manipulation.",
        ],
        "I. Introduction Continued": [
            "The project therefore needs to be judged on two axes. The first is machine-learning quality: data preparation, model training, validation accuracy, metrics, and reproducibility. The second is security quality: attack implementation, threat-model clarity, trigger behavior, defense evaluation, calibration, and evidence that mitigation reduces malicious behavior.",
            "This report uses the implemented codebase as its primary evidence source. Empty or sparse documentation files are treated as documentation gaps rather than as hidden knowledge. Where the source code and generated artifacts provide clear numbers, those numbers are reported directly; where the repository does not provide enough evidence, the limitation is stated explicitly.",
        ],
        "II. Literature Review": [
            "BadNets introduced the practical idea that a model can learn a hidden association between a trigger and a target label while preserving normal accuracy. In benchmark terms, this means clean accuracy is necessary but insufficient: attack success rate must be measured separately using triggered inputs.",
            "The attack families implemented in BTSecBench_v2 reflect important directions in the literature. Patch attacks test obvious localized features, blend attacks test distributed visual contamination, signal attacks test periodic perturbations, and warping attacks test geometric transformations that may be harder to detect by simple pixel inspection.",
        ],
        "II. Literature Review Continued": [
            "STRIP belongs to a class of runtime detection methods. It does not attempt to repair the model. Instead, it attempts to recognize suspicious inputs by examining prediction stability under perturbations. This is attractive because it can be deployed around an existing model, but its usefulness depends heavily on false positive rate, calibration data, and operational tolerance for rejected inputs.",
            "Fine-tuning belongs to mitigation rather than detection. It can sometimes reduce reliance on trigger shortcuts by reinforcing clean decision boundaries. However, if the trigger association is deeply embedded or if clean fine-tuning is too short, the backdoor can remain active while clean validation accuracy looks excellent.",
        ],
        "II. Literature Review Summary": [
            "A strong empirical paper in this area usually reports a matrix of results rather than a single experiment. The expected matrix includes model family, attack type, poison rate, target class, random seed, clean accuracy, ASR, defense method, post-defense clean accuracy, post-defense ASR, and computational cost.",
            "BTSecBench_v2 already implements enough components to move toward this standard. The missing piece is not a new conceptual architecture but a systematic experiment runner that exercises the existing components consistently and records every run with complete metadata.",
        ],
        "III. Problem Statement and Objectives": [
            "The formal problem can be stated as follows: given a classifier trained for 43-class traffic sign recognition, evaluate whether maliciously triggered inputs can force predictions toward a target class and whether available defenses can detect or reduce that behavior without sacrificing clean performance.",
            "The benchmark must also support fair comparison. Fairness here means using the same dataset splits, preprocessing pipeline, metric definitions, seed policy, and reporting schema across attacks and defenses. Without these controls, observed differences may come from experiment setup rather than from the underlying method.",
        ],
        "IV. Dataset and Exploratory Data Analysis": [
            "The observed image-size variation is important because resizing can change trigger geometry. A 4 by 4 patch on a 64 by 64 image has a fixed relative footprint, but if triggers are applied before resizing, their visual footprint can change. The benchmark should therefore document whether triggers are applied before or after the final transform pipeline.",
            "Class imbalance should also be considered. In traffic sign datasets, some signs occur more frequently than others. If the target class is common, attack-success interpretation may differ from a rare target class. Per-class support and per-class ASR would make the benchmark more informative.",
        ],
        "IV. Dataset and EDA Continued": [
            "The existing statistics files are a good starting point because they make the dataset measurable. Additional plots could include representative samples per class, train/validation class balance, pixel mean and standard deviation by channel, and before/after transform examples.",
            "For security analysis, EDA should include trigger visibility studies. Each attack should generate example clean and poisoned images after preprocessing so that readers can see what the model actually receives. This is especially important for SIG and WaNet-style triggers.",
        ],
        "IV. Dataset Pipeline": [
            "The dataset abstraction is simple and appropriate. Returning a dictionary rather than a tuple allows later components to carry image, label, and path together. This design is helpful for explainability and error analysis because misclassified or suspicious samples can be traced back to source files.",
            "One improvement would be to store dataset metadata alongside every benchmark result. The metadata should include split file names, split hashes, image root, transform mode, normalization values, and sample counts. These fields would make later audits much easier.",
        ],
        "V. Methodology": [
            "The methodology is best described as componentized experimental orchestration. Each subsystem has a narrow responsibility: data loading prepares samples, models produce logits, attacks modify inputs and labels, the trainer optimizes parameters, evaluators compute metrics, defenses detect or mitigate backdoors, and reporting utilities persist evidence.",
            "This separation is important because backdoor experiments are prone to confounding. If attack code, training code, and evaluation code are tightly coupled, it becomes hard to know whether a result is caused by the attack method or by implementation details. BTSecBench_v2 mostly avoids this by using factories and reusable classes.",
        ],
        "V. Methodology: Training and Evaluation": [
            "Gradient clipping is a useful stability choice, especially when experimenting with multiple models and potentially noisy poisoned samples. The training loop also records throughput and time, which matter because defense methods such as STRIP can multiply inference cost.",
            "The evaluator should remain the single source of truth for metric definitions. If multiple scripts compute metrics independently, small inconsistencies can appear. A mature benchmark should route all accuracy, precision, recall, F1, balanced accuracy, MCC, and confusion-matrix calculations through shared tested utilities.",
        ],
        "V. Methodology: Attack Evaluation": [
            "Attack success rate should be defined carefully. For targeted attacks, ASR usually measures the fraction of triggered non-target samples classified as the attacker's target class. The benchmark should document whether target-class source samples are excluded, how labels are remapped, and whether clean accuracy on untriggered samples is measured in the same run.",
            "Because the current BadNets artifact uses poison rate 1.0 for validation stress, it should be distinguished from training poison rate. A validation poison rate of 1.0 is useful for measuring triggered behavior, while a training poison rate controls how much compromised data the model sees during learning.",
        ],
        "V. Methodology: Defense Evaluation": [
            "Defense evaluation should include both detection metrics and model-utility metrics. For STRIP, the essential values are threshold, precision, recall, F1, false positive rate, false negative rate, AUC, and confusion matrix. For fine-tuning, the essential values are clean accuracy, post-defense ASR, training cost, and any degradation in calibration or class balance.",
            "A fair defense comparison should use a separate calibration split when thresholds are learned. Selecting a threshold and reporting performance on the same data can produce optimistic results. This is particularly relevant for STRIP because the chosen threshold directly controls the false positive and false negative trade-off.",
        ],
        "V. Methodology: Reproducibility": [
            "Reproducibility is not only about setting a seed. It also requires stable dependencies, documented commands, immutable split files, stored configuration, versioned outputs, and enough metadata to reconstruct the experiment environment. BTSecBench_v2 has several of these pieces but should make them mandatory for every benchmark run.",
            "A useful next step is a manifest file for each experiment. The manifest should include model name, attack name, defense name, target class, poison rate, batch size, learning rate, epochs, seed, checkpoint path, dataset split hashes, code commit, hardware, and timestamp.",
        ],
        "VI. Deep Learning Model Comparison": [
            "The four supported backbones cover a practical spectrum. The CNN baseline is simple and transparent; ResNet18 is a widely used residual architecture; MobileNetV3 is designed for efficiency; and EfficientNet-B0 balances compound scaling with relatively compact parameter count.",
            "For traffic sign recognition, input resolution is small compared with many ImageNet settings. This means transfer-learning backbones may be powerful, but their behavior should be validated carefully at 64 by 64 resolution. Differences in receptive field, normalization expectations, and classifier replacement can affect both clean accuracy and backdoor sensitivity.",
        ],
        "VI. Model Comparison Discussion": [
            "Runtime results should be interpreted in the context of defense cost. A model that is only moderately slower under ordinary inference may become significantly more expensive under STRIP because each input is evaluated many times. Therefore, defense-adjusted throughput should be included in future model comparisons.",
            "The model benchmark also provides a sanity check for implementation quality. Matching output shape `(batch, 43)` across all backbones confirms that the registry and classifier heads are wired correctly for the GTSRB task.",
        ],
        "VII. Backdoor Attack Implementation": [
            "The attack base class is one of the strongest design choices in the repository. By standardizing `poison_sample`, `poison_batch`, and `__call__`, it allows datasets, scripts, and benchmark runners to use attacks interchangeably. This reduces duplication and makes it easier to add future attacks.",
            "Future attack implementations should include a `describe()` or metadata method that exports trigger size, color, alpha, frequency, warp strength, target class, and poison rate. That metadata should be saved into benchmark JSON so that figures and tables can be interpreted without reading source code.",
        ],
        "VII. BadNets and Blend Attacks": [
            "BadNets is useful for debugging because the trigger is explicit and localized. If ASR is unexpectedly low, researchers can inspect whether the patch is visible after transforms, whether target labels were applied correctly, and whether the poisoned model was trained long enough to learn the shortcut.",
            "Blend attacks are more subtle because the perturbation is distributed over the whole image. They may interact differently with normalization and data augmentation. The benchmark should therefore report both visual examples and numeric metrics for each alpha value.",
        ],
        "VII. SIG and WaNet Attacks": [
            "SIG and WaNet increase benchmark diversity. SIG introduces a periodic intensity pattern, while WaNet changes spatial geometry without necessarily adding an obvious patch. These attacks are important because defenses tuned for visible patches may fail against non-patch triggers.",
            "The WaNet implementation uses random displacement fields, so seed handling is especially important. The benchmark should verify that repeated runs with the same seed produce repeatable warps or should explicitly document stochastic behavior if randomness is intended at evaluation time.",
        ],
        "VIII. STRIP Detection": [
            "STRIP's central assumption is that a triggered sample will remain strongly associated with the target class even when blended with unrelated clean images. Clean samples should produce more variable predictions under perturbation, yielding higher entropy. This difference becomes the detection signal.",
            "The effectiveness of this signal depends on perturbation count, blend alpha, model calibration, attack type, and threshold policy. A single threshold may not transfer across attacks or backbones. Future reports should include threshold-sensitivity curves and AUC values by attack.",
        ],
        "VIII. STRIP Results": [
            "The reported high recall means STRIP is catching most poisoned samples in the tested setting. However, the high false positive rate means many clean samples would be blocked or escalated. In a real deployment, this trade-off may be unacceptable unless a second-stage review or fallback classifier is available.",
            "The confusion matrix should be discussed alongside system requirements. If missing a poisoned input is extremely costly, a high-recall operating point may be justified. If availability and clean throughput are important, the threshold must be relaxed or the detector must be combined with another signal.",
        ],
        "IX. Fine-Tuning Defense": [
            "Fine-tuning is attractive because it is simple to implement and easy to explain. It uses clean data to adjust model parameters after a suspected compromise. In BTSecBench_v2, this is implemented through the same trainer and evaluator used elsewhere, which reduces the risk of inconsistent training behavior.",
            "The limitation is that fine-tuning may preserve the backdoor if the clean data do not contradict the trigger association. The model can remain accurate on clean validation samples while still responding to a trigger. Therefore, post-defense triggered evaluation is mandatory for a complete claim.",
        ],
        "IX. Fine-Tuning Results": [
            "The fine-tuning history suggests stable clean validation behavior across epochs. Validation loss decreases and accuracy remains near 99.95 percent. This indicates the model has strong clean classification capacity after fine-tuning.",
            "For a stronger report, this section should add before/after ASR, per-class clean accuracy, per-class recall, and trigger-specific examples. If fine-tuning reduces ASR while preserving clean accuracy, the defense claim becomes much stronger.",
        ],
        "X. Experimental Results and Benchmark": [
            "The current benchmark table is valuable because it brings attack, detection, and defense evidence into one place. It shows that the project can move beyond isolated scripts and produce reportable artifacts. However, the table should be expanded into a full experiment matrix before broad conclusions are drawn.",
            "The figure outputs make the benchmark easier to interpret. Accuracy comparison, ASR comparison, STRIP metrics, ROC curve, entropy histogram, entropy distribution, confusion matrix, and benchmark summary images are present in the reports directory. Future figures should embed run identifiers and trace back to exact JSON inputs.",
        ],
        "XI. Error Analysis and Discussion": [
            "A false positive in STRIP means a clean traffic sign is treated as suspicious. Depending on deployment policy, this could cause abstention, manual review, fallback inference, or delayed response. Thus, false positives are not merely statistical errors; they have system-level consequences.",
            "A false negative means a poisoned input passes through the detector. In a safety-relevant setting, false negatives can be more dangerous than false positives. The correct operating point depends on whether the system prioritizes safety, availability, or a balanced trade-off.",
        ],
        "XI. Discussion Continued": [
            "The documentation gap is important because reproducible security work must be executable by someone other than the original author. A good README should include environment setup, dataset download, split generation, model training, backdoor training, attack evaluation, STRIP evaluation, fine-tuning defense, and report generation commands.",
            "The test gap is equally important. Backdoor benchmarks involve many subtle transformations, and a small bug in label remapping, tensor conversion, normalization, or thresholding can change the conclusion. Automated tests should protect these assumptions.",
        ],
        "XII. Conclusion and Future Work": [
            "The most important future work is systematic experimentation. The project should run all supported backbones against all implemented attacks under multiple seeds and poison rates. Each defense should be evaluated on the same matrix, with clean accuracy and ASR reported before and after defense.",
            "The second priority is artifact quality. Documentation, packaging metadata, test coverage, run manifests, and reproducible commands will make the project easier to evaluate, reuse, and cite. These improvements are engineering work, but they directly affect research credibility.",
        ],
        "References": [
            "Additional reference categories that would strengthen a final academic version include recent surveys on backdoor learning, adaptive backdoor defenses, calibration under distribution shift, and robust evaluation methodology. These should be added after the benchmark matrix is finalized so that citations match the exact experiments reported.",
        ],
    }

    for title_text, paragraphs, bullets, table, figure in pages:
        expanded_paragraphs = list(paragraphs or []) + section_expansions.get(title_text, [])
        add_page(doc, title_text, expanded_paragraphs, bullets, table, figure)

    doc.save(OUT)


if __name__ == "__main__":
    build()
    print(OUT)
